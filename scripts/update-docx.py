"""Rebuild the Claude Skills Suite Reference.docx with current tools and skills."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches

DOCX_PATH = r'C:\Dev\claude-skills-suite\Claude Skills Suite Reference.docx'

doc = Document(DOCX_PATH)

# === Clear existing content ===
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
for t in list(doc.tables):
    t._element.getparent().remove(t._element)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = doc.styles['Light Shading Accent 1']
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            if ci == 0 and val.startswith('/'):
                run = p.add_run(val)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.bold = True
            else:
                run = p.add_run(val)
                run.font.size = Pt(10)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Inches(w)
    return table


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ================================================================
# TITLE
# ================================================================
doc.add_paragraph(
    'Claude Skills Suite \u2014 Tools & Skills Reference', style='Title'
)

p = doc.add_paragraph()
run = p.add_run(
    '44 skills + 35 MCP gateway modules + native Claude Code tools.'
)
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Updated: 2026-03-30  |  Repo: mrschedler/claude-skills-suite')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ================================================================
# NATIVE CLAUDE CODE TOOLS
# ================================================================
doc.add_heading('Native Claude Code Tools', level=1)
add_table(doc, ['Tool', 'Purpose'], [
    ['Read',        'Read files by absolute path (text, images, PDFs, notebooks)'],
    ['Write',       'Create or overwrite files'],
    ['Edit',        'Exact string replacement in existing files'],
    ['Glob',        'Fast file pattern matching (e.g. **/*.ts)'],
    ['Grep',        'Regex content search powered by ripgrep'],
    ['Bash',        'Shell command execution (Git Bash on Windows)'],
    ['Agent',       'Spawn specialized subagents (Explore, Plan, general-purpose)'],
    ['WebFetch',    'Fetch and process URL content'],
    ['WebSearch',   'Web search'],
    ['TaskCreate / TaskUpdate', 'Break work into tracked steps within a session'],
    ['Skill',       'Invoke a registered slash-command skill'],
    ['NotebookEdit', 'Edit Jupyter notebook cells'],
], col_widths=[2.0, 4.5])

# ================================================================
# MCP GATEWAY MODULES
# ================================================================
doc.add_heading('MCP Gateway (35 modules, 300+ tools)', level=1)
add_subtitle(doc,
    'Endpoint: https://mcp.epiphanyco.com/mcp \u2014 '
    'each module exposes _call (execute) and _list (discover) endpoints.')

# -- Knowledge & Memory --
doc.add_heading('Knowledge & Memory', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['memory',   'Qdrant semantic memory \u2014 decisions, reasoning, gotchas, preferences, session history'],
    ['graph',    'Neo4j knowledge graph \u2014 structural relationships, dependencies, entity connections'],
    ['obsidian', 'Obsidian vault indexing and semantic search'],
], col_widths=[1.5, 5.0])

# -- Project & Task Management --
doc.add_heading('Project & Task Management', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['project',   'Project pipeline lifecycle \u2014 sprints, tasks, questions, events'],
    ['task',      'Universal work tracking \u2014 punch lists, diagnostics, infrastructure, research, stories'],
    ['blueprint', 'Agent blueprints \u2014 modular automation specs with guardrails'],
    ['workspace', 'Project workspace orchestration \u2014 repos, sessions, designs, artifacts'],
], col_widths=[1.5, 5.0])

# -- Data Storage --
doc.add_heading('Data Storage & Querying', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['pg',       'PostgreSQL database tools \u2014 query, manage tables'],
    ['mongodb',  'MongoDB \u2014 databases, collections, documents, indexes, aggregation'],
    ['redis',    'Shared Redis cache/message broker \u2014 key ops, pub/sub, server stats'],
    ['rabbitmq', 'RabbitMQ message queue \u2014 exchanges, bindings, connections'],
], col_widths=[1.5, 5.0])

# -- Infrastructure --
doc.add_heading('Infrastructure & DevOps', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['docker',     'Docker container management \u2014 inspect, run, stop, restart'],
    ['recipe',     'Container recipes \u2014 reusable templates stored in Vault'],
    ['traefik',    'Traefik reverse proxy inspection (read-only) \u2014 routers, services, middlewares'],
    ['vault',      'HashiCorp Vault secrets management'],
    ['prometheus', 'Prometheus metrics and monitoring queries'],
    ['uptime',     'Uptime Kuma service availability monitoring'],
], col_widths=[1.5, 5.0])

# -- Automation & AI --
doc.add_heading('Automation & AI', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['n8n',       'Workflow automation \u2014 list, create, update, execute, delete workflows'],
    ['dify',      'Self-hosted Dify AI app builder \u2014 apps, models, tools, knowledge bases'],
    ['openhands', 'Self-hosted OpenHands autonomous coding agent'],
    ['ai',        'OpenRouter AI \u2014 call other LLMs (Gemini, GPT, Mistral, etc.)'],
    ['context7',  'Up-to-date documentation for code libraries and frameworks'],
], col_widths=[1.5, 5.0])

# -- External Integrations --
doc.add_heading('External Integrations', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['github',     'GitHub \u2014 repos, issues, PRs, checks, releases'],
    ['browser',    'Browser automation \u2014 Playwright-based web interactions'],
    ['figma',      'Figma design files \u2014 read files, extract components/styles, render images'],
    ['mattermost', 'Team chat at chat.epiphanyco.com (LAN-only, backup notifications only)'],
    ['telegram',   'Telegram bot \u2014 send messages, photos, inline keyboards, webhooks'],
    ['contacts',   'Email intelligence \u2014 contact dossiers, email search, business context'],
], col_widths=[1.5, 5.0])

# -- Coordination --
doc.add_heading('Coordination & Configuration', level=1)
add_table(doc, ['Module', 'Purpose'], [
    ['gateway',      'Gateway orchestration \u2014 includes rehydrate for session startup'],
    ['coordination', 'Inter-agent session registration, coordination signals, resource locks'],
    ['interagent',   'Inter-agent messaging \u2014 inbox, claim, complete across machines'],
    ['activity',     'Activity log querying \u2014 track all tool calls across sessions'],
    ['pref',         'User preferences (legacy; migrated to Qdrant category: preference)'],
], col_widths=[1.5, 5.0])

# ================================================================
# LOCAL MCP SERVERS
# ================================================================
doc.add_heading('Local MCP Servers', level=1)

doc.add_heading('Homelab (mcp__homelab__*)', level=1)
add_table(doc, ['Tool', 'Purpose'], [
    ['list_workflows / get_workflow',         'Inspect n8n workflows'],
    ['create / update / delete_workflow',     'Manage n8n workflows'],
    ['activate / deactivate_workflow',        'Toggle n8n workflow state'],
    ['list_executions / get_execution',       'View n8n execution history'],
    ['query_mariadb / list_mariadb_tables',   'Tour database (MariaDB)'],
    ['query_postgres / list_postgres_tables', 'Mailmine database (PostgreSQL)'],
    ['describe_table',                        'Schema inspection for either DB'],
    ['list_containers / container_logs',      'Docker container inspection'],
    ['container_restart / stop / start',      'Docker container lifecycle'],
    ['ssh_command / ssh_device',              'SSH to Unraid, Jetson, Pi-106, Pi-105'],
], col_widths=[2.8, 3.7])

doc.add_heading('Playwright', level=1)
add_subtitle(doc,
    'Browser automation via @playwright/mcp@latest. Used by the /browser-review skill.')

# ================================================================
# GMAIL INTEGRATION
# ================================================================
doc.add_heading('Gmail Integration', level=1)
add_table(doc, ['Tool', 'Purpose'], [
    ['gmail_get_profile',      'User profile and mailbox stats'],
    ['gmail_search_messages',  'Search with full Gmail query syntax'],
    ['gmail_read_message',     'Fetch complete message content'],
    ['gmail_read_thread',      'Fetch entire conversation thread'],
    ['gmail_list_drafts',      'List unsent drafts with pagination'],
    ['gmail_list_labels',      'List system and user-created labels'],
    ['gmail_create_draft',     'Create new draft (plain text or HTML, reply-to-thread)'],
], col_widths=[2.5, 4.0])

# ================================================================
# SKILLS SUITE
# ================================================================
doc.add_heading('Skills Suite (44 skills)', level=1)
add_subtitle(doc,
    'Invoked via /skill-name in Claude Code. Each skill lives in skills/<name>/SKILL.md.')

# -- Configuration --
doc.add_heading('Configuration', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/update-config',    'Configure Claude Code settings, hooks, permissions, and env vars'],
    ['/keybindings-help', 'Customize keyboard shortcuts and rebind keys'],
])

# -- Development --
doc.add_heading('Development', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/feature-dev',      'Unified feature development \u2014 routes by complexity: simple (do it), medium (light plan), complex (Ralph mode with PRD + artifact DB)'],
    ['/claude-api',       'Build apps with the Claude API or Anthropic SDK'],
    ['/ui-design',        'Generate UI components and pages following the project design system'],
    ['/simplify',         'Review changed code for reuse, quality, and efficiency, then fix issues'],
])

# -- Planning --
doc.add_heading('Planning', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/quick-plan',        'Lightweight in-session planning with phases and acceptance criteria'],
    ['/build-plan',        'Formal project-plan.md with phases, milestones, and parallelizable work units'],
    ['/sub-project',       'Create an isolated sub-project workspace within a parent project'],
    ['/sub-project-merge', 'Merge a completed sub-project back into its parent'],
])

# -- Project Lifecycle --
doc.add_heading('Project Lifecycle', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/meta-init',         'Full new-project workflow: scaffold \u2192 interview \u2192 context \u2192 build plan'],
    ['/meta-join',         'Join/onboard to an existing project (full onboard or quick catch-up)'],
    ['/project-organize',  'Organize any project with GROUNDING.md, notebook, and clean structure'],
    ['/project-context',   'Write project-context.md \u2014 a comprehensive handoff document'],
    ['/notebook-init',     "Create an engineering or inventor\u2019s notebook in a project"],
    ['/evolve',            'Update project-context.md and project-plan.md to reflect current truth'],
    ['/github-pull',       'Pull latest changes from remote'],
    ['/github-sync',       'Commit and push changes to GitHub'],
])

# -- Research --
doc.add_heading('Research', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/claude-light-research',  'Lightweight research with artifact DB storage \u2014 no subagents, no debate, everyday use'],
    ['/claude-deep-research',  'Claude-only deep research with steelman debate (~15 workers, convergence scoring)'],
])

# -- Review Lenses --
doc.add_heading('Review Lenses', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/meta-review',            'Run multiple review lenses in parallel and synthesize findings'],
    ['/security-review',        'Security audit: dependencies, auth, secrets, input validation, supply chain'],
    ['/completeness-review',    'Scan for stubs, TODOs, placeholders, empty bodies, unfinished code'],
    ['/test-review',            'Evaluate test coverage, quality, gaps, and AI tendencies to skip tests'],
    ['/drift-review',           'Compare code against documentation to find drift'],
    ['/refactor-review',        'Code quality pass: over-engineering, duplication, bloat, truncated code'],
    ['/perf-review',            'N+1 queries, missing indexes, memory leaks, O(n\u00b2) loops, caching gaps'],
    ['/log-review',             'Audit logging and observability: silent catches, missing context, no trace IDs'],
    ['/integration-review',     'Dead wiring, missing config/env, incomplete teardown, unbundled assets'],
    ['/doc-audit',              'Documentation quality, completeness, accuracy, and doc-code drift'],
    ['/dep-audit',              'Dependency health: CVEs, outdated versions, license conflicts, abandoned packages'],
    ['/ui-review',              'UI anti-patterns, token violations, a11y failures, inconsistency'],
    ['/browser-review',         'Visual QA via browser MCP tools (Playwright/browser-use)'],
    ['/breaking-change-review', 'Detect breaking API, dependency, and schema changes before they ship'],
    ['/compliance-review',      'Check code against documented rules (CLAUDE.md, cross-cutting-rules.md)'],
    ['/counter-review',         'Adversarial red-team review: architecture, abuse cases, attack chains, failures'],
    ['/review-fix',             'Implement fixes from review findings with user approval'],
])

# -- Meta / Orchestration --
doc.add_heading('Meta / Orchestration', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/meta-production',    'Scored production readiness assessment across 12 dimensions (READY / NOT READY)'],
    ['/meta-context-save',  'Save session state and optionally commit + clear context'],
])

# -- Skill Management --
doc.add_heading('Skill Management', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/skill-forge',  'Create or edit skills \u2014 scaffolds directory, writes SKILL.md from template'],
    ['/skill-doctor', 'Self-diagnostic for the skill suite \u2014 run after install or when skills fail'],
])

# -- Automation --
doc.add_heading('Automation', level=1)
add_table(doc, ['Skill', 'Purpose'], [
    ['/loop',     'Run a prompt or slash command on a recurring interval (e.g., /loop 5m /foo)'],
    ['/schedule', 'Create, update, list, or run scheduled remote agents on a cron schedule'],
])

# -- Internal Skills --
doc.add_heading('Internal Skills (not user-invocable)', level=1)
add_subtitle(doc, 'These are called by other skills, not invoked directly.')

add_table(doc, ['Skill', 'Category', 'Purpose'], [
    ['/claude-deep-research-execute',  'Research',       'Opus subagent for Claude-only deep research'],
    ['/deploy-gateway',                'Infrastructure', 'Deploy, rebuild, or restart the MCP gateway container'],
    ['/infra-health',                  'Infrastructure', 'Check service health \u2014 containers, endpoints, monitoring'],
    ['/init-db',                       'Database',       'Initialize artifact store (SQLite+FTS5) \u2014 idempotent'],
    ['/test-gen',                      'Development',    'Generate tests from test-review findings'],
    ['/log-gen',                       'Development',    'Generate logging instrumentation from log-review findings'],
    ['/project-questions',             'Project Setup',  'Deep-dive interview to surface assumptions and constraints'],
])

# ================================================================
# QUICK REFERENCE
# ================================================================
doc.add_heading('Quick Reference: Common Workflows', level=1)
add_table(doc, ['I want to...', 'Use'], [
    ['Start a new project',          '/meta-init'],
    ['Join an existing project',     '/meta-join'],
    ['Build a feature',              '/feature-dev'],
    ['Plan before building',         '/build-plan or /quick-plan'],
    ['Run a comprehensive review',   '/meta-review'],
    ['Fix review findings',          '/review-fix'],
    ['Check if we can ship',         '/meta-production'],
    ['Research a topic',             '/claude-light-research'],
    ['Deep research a topic',        '/claude-deep-research'],
    ['Commit and push',              '/github-sync'],
    ['Prepare a release',            'gh release create (no skill needed)'],
    ['Create a new skill',           '/skill-forge'],
    ['Diagnose skill issues',        '/skill-doctor'],
    ['Save session and wrap up',     '/meta-context-save'],
    ['Search past decisions',        'memory_call > search (Qdrant)'],
    ['Find entity relationships',    'graph_call > query (Neo4j)'],
    ['Check project status',         'project_call > get_project'],
    ['Manage secrets',               'vault_call'],
    ['Monitor services',             '/infra-health or uptime_call'],
])

# Save
doc.save(DOCX_PATH)
print('Document saved successfully.')
